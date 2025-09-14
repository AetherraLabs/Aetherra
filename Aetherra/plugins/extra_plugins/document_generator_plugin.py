"""
Document Generator Plugin - Professional Document Creation and Template System
Author: Aetherra Plugin System
Version: 1.0.0

This plugin provides comprehensive document generation capabilities including:
- Resume/CV builder with professional templates
- Meeting notes and report generation
- Markdown to PDF/Word conversion
- Template management system
- AI-enhanced content suggestions
- Multi-format export (PDF, DOCX, HTML, Markdown)
"""

import json
import logging
import os
import sys
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    import markdown
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    from docx.shared import Inches, Pt

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Set up logging
logger = logging.getLogger(__name__)


@dataclass
class DocumentTemplate:
    """Document template structure."""

    id: str
    name: str
    category: str
    description: str
    fields: list[dict[str, Any]]
    content_template: str
    style_config: dict[str, Any]
    output_formats: list[str]
    created_date: str
    author: str


@dataclass
class DocumentData:
    """Document data structure."""

    template_id: str
    title: str
    content: dict[str, Any]
    metadata: dict[str, Any]
    created_date: str
    modified_date: str
    output_format: str


class TemplateEngine:
    """Template processing and rendering engine."""

    def __init__(self):
        self.templates = {}
        self.load_built_in_templates()

    def load_built_in_templates(self):
        """Load built-in document templates."""
        # Resume/CV Template
        resume_template = DocumentTemplate(
            id="professional_resume",
            name="Professional Resume",
            category="career",
            description="Modern professional resume template with sections for experience, education, and skills",
            fields=[
                {
                    "name": "personal_info",
                    "type": "object",
                    "required": True,
                    "fields": [
                        {"name": "full_name", "type": "string", "required": True},
                        {"name": "email", "type": "email", "required": True},
                        {"name": "phone", "type": "string", "required": False},
                        {"name": "location", "type": "string", "required": False},
                        {"name": "website", "type": "url", "required": False},
                        {"name": "linkedin", "type": "url", "required": False},
                    ],
                },
                {"name": "professional_summary", "type": "text", "required": True},
                {
                    "name": "experience",
                    "type": "array",
                    "required": True,
                    "item_fields": [
                        {"name": "company", "type": "string", "required": True},
                        {"name": "position", "type": "string", "required": True},
                        {"name": "start_date", "type": "date", "required": True},
                        {"name": "end_date", "type": "date", "required": False},
                        {"name": "description", "type": "text", "required": True},
                        {"name": "achievements", "type": "array", "required": False},
                    ],
                },
                {
                    "name": "education",
                    "type": "array",
                    "required": True,
                    "item_fields": [
                        {"name": "institution", "type": "string", "required": True},
                        {"name": "degree", "type": "string", "required": True},
                        {"name": "field", "type": "string", "required": True},
                        {"name": "graduation_date", "type": "date", "required": True},
                        {"name": "gpa", "type": "number", "required": False},
                    ],
                },
                {"name": "skills", "type": "array", "required": True},
                {"name": "certifications", "type": "array", "required": False},
                {"name": "projects", "type": "array", "required": False},
            ],
            content_template="""
# {personal_info.full_name}

**Email:** {personal_info.email} | **Phone:** {personal_info.phone}
**Location:** {personal_info.location} | **Website:** {personal_info.website}
**LinkedIn:** {personal_info.linkedin}

## Professional Summary
{professional_summary}

## Work Experience
{%- for exp in experience %}
### {exp.position} at {exp.company}
**{exp.start_date} - {exp.end_date or 'Present'}**

{exp.description}

{%- if exp.achievements %}
**Key Achievements:**
{%- for achievement in exp.achievements %}
- {achievement}
{%- endfor %}
{%- endif %}
{%- endfor %}

## Education
{%- for edu in education %}
### {edu.degree} in {edu.field}
**{edu.institution}** - {edu.graduation_date}
{%- if edu.gpa %}
GPA: {edu.gpa}
{%- endif %}
{%- endfor %}

## Skills
{%- for skill in skills %}
- {skill}
{%- endfor %}

{%- if certifications %}
## Certifications
{%- for cert in certifications %}
- {cert}
{%- endfor %}
{%- endif %}

{%- if projects %}
## Notable Projects
{%- for project in projects %}
- {project}
{%- endfor %}
{%- endif %}
            """,
            style_config={
                "font_family": "Helvetica",
                "font_size": 11,
                "line_spacing": 1.2,
                "margins": {"top": 72, "bottom": 72, "left": 72, "right": 72},
                "colors": {
                    "primary": "#2C3E50",
                    "secondary": "#34495E",
                    "accent": "#3498DB",
                },
            },
            output_formats=["pdf", "docx", "html", "markdown"],
            created_date=datetime.now().isoformat(),
            author="Aetherra Document Generator",
        )

        # Meeting Notes Template
        meeting_template = DocumentTemplate(
            id="meeting_notes",
            name="Meeting Notes",
            category="business",
            description="Structured meeting notes template with agenda, attendees, and action items",
            fields=[
                {"name": "meeting_title", "type": "string", "required": True},
                {"name": "date", "type": "date", "required": True},
                {"name": "time", "type": "time", "required": True},
                {"name": "duration", "type": "string", "required": False},
                {"name": "location", "type": "string", "required": False},
                {"name": "organizer", "type": "string", "required": True},
                {"name": "attendees", "type": "array", "required": True},
                {"name": "agenda_items", "type": "array", "required": True},
                {"name": "discussions", "type": "text", "required": True},
                {"name": "decisions", "type": "array", "required": False},
                {
                    "name": "action_items",
                    "type": "array",
                    "required": True,
                    "item_fields": [
                        {"name": "task", "type": "string", "required": True},
                        {"name": "assignee", "type": "string", "required": True},
                        {"name": "due_date", "type": "date", "required": False},
                        {"name": "priority", "type": "string", "required": False},
                    ],
                },
                {
                    "name": "next_meeting",
                    "type": "object",
                    "required": False,
                    "fields": [
                        {"name": "date", "type": "date", "required": True},
                        {"name": "time", "type": "time", "required": True},
                        {"name": "agenda", "type": "string", "required": False},
                    ],
                },
            ],
            content_template="""
# Meeting Notes: {meeting_title}

**Date:** {date}
**Time:** {time}
{%- if duration %}
**Duration:** {duration}
{%- endif %}
{%- if location %}
**Location:** {location}
{%- endif %}
**Organizer:** {organizer}

## Attendees
{%- for attendee in attendees %}
- {attendee}
{%- endfor %}

## Agenda
{%- for item in agenda_items %}
- {item}
{%- endfor %}

## Discussion Summary
{discussions}

{%- if decisions %}
## Key Decisions
{%- for decision in decisions %}
- {decision}
{%- endfor %}
{%- endif %}

## Action Items
{%- for action in action_items %}
- **{action.task}**
  - Assignee: {action.assignee}
  {%- if action.due_date %}
  - Due Date: {action.due_date}
  {%- endif %}
  {%- if action.priority %}
  - Priority: {action.priority}
  {%- endif %}
{%- endfor %}

{%- if next_meeting %}
## Next Meeting
**Date:** {next_meeting.date}
**Time:** {next_meeting.time}
{%- if next_meeting.agenda %}
**Agenda:** {next_meeting.agenda}
{%- endif %}
{%- endif %}
            """,
            style_config={
                "font_family": "Arial",
                "font_size": 12,
                "line_spacing": 1.4,
                "margins": {"top": 72, "bottom": 72, "left": 72, "right": 72},
                "colors": {
                    "primary": "#2C3E50",
                    "secondary": "#7F8C8D",
                    "accent": "#E74C3C",
                },
            },
            output_formats=["pdf", "docx", "html", "markdown"],
            created_date=datetime.now().isoformat(),
            author="Aetherra Document Generator",
        )

        # Technical Report Template
        report_template = DocumentTemplate(
            id="technical_report",
            name="Technical Report",
            category="technical",
            description="Comprehensive technical report template for projects and analysis",
            fields=[
                {"name": "title", "type": "string", "required": True},
                {"name": "author", "type": "string", "required": True},
                {"name": "date", "type": "date", "required": True},
                {"name": "version", "type": "string", "required": False},
                {"name": "executive_summary", "type": "text", "required": True},
                {"name": "introduction", "type": "text", "required": True},
                {"name": "methodology", "type": "text", "required": True},
                {"name": "findings", "type": "text", "required": True},
                {"name": "recommendations", "type": "array", "required": True},
                {"name": "conclusion", "type": "text", "required": True},
                {"name": "references", "type": "array", "required": False},
                {"name": "appendices", "type": "array", "required": False},
            ],
            content_template="""
# {title}

**Author:** {author}
**Date:** {date}
{%- if version %}
**Version:** {version}
{%- endif %}

## Executive Summary
{executive_summary}

## 1. Introduction
{introduction}

## 2. Methodology
{methodology}

## 3. Findings
{findings}

## 4. Recommendations
{%- for rec in recommendations %}
{loop.index}. {rec}
{%- endfor %}

## 5. Conclusion
{conclusion}

{%- if references %}
## References
{%- for ref in references %}
{loop.index}. {ref}
{%- endfor %}
{%- endif %}

{%- if appendices %}
## Appendices
{%- for appendix in appendices %}
### Appendix {loop.index}
{appendix}
{%- endfor %}
{%- endif %}
            """,
            style_config={
                "font_family": "Times New Roman",
                "font_size": 12,
                "line_spacing": 1.5,
                "margins": {"top": 90, "bottom": 90, "left": 90, "right": 90},
                "colors": {
                    "primary": "#000000",
                    "secondary": "#333333",
                    "accent": "#1E88E5",
                },
            },
            output_formats=["pdf", "docx", "html", "markdown"],
            created_date=datetime.now().isoformat(),
            author="Aetherra Document Generator",
        )

        # Store templates
        self.templates = {
            "professional_resume": resume_template,
            "meeting_notes": meeting_template,
            "technical_report": report_template,
        }

    def get_template(self, template_id: str) -> DocumentTemplate | None:
        """Get a template by ID."""
        return self.templates.get(template_id)

    def list_templates(self, category: str | None = None) -> list[DocumentTemplate]:
        """List available templates, optionally filtered by category."""
        templates = list(self.templates.values())
        if category:
            templates = [t for t in templates if t.category == category]
        return templates

    def render_template(self, template_id: str, data: dict[str, Any]) -> str:
        """Render a template with the provided data."""
        template = self.get_template(template_id)
        if not template:
            raise ValueError(f"Template {template_id} not found")

        try:
            # Simple template rendering (in a real implementation, use Jinja2)
            content = template.content_template

            # Basic variable substitution
            content = self._substitute_variables(content, data)

            return content
        except Exception as e:
            logger.error(f"Template rendering failed: {e}")
            raise

    def _substitute_variables(self, content: str, data: dict[str, Any]) -> str:
        """Simple variable substitution for templates."""
        # This is a simplified implementation
        # In production, use a proper template engine like Jinja2

        def replace_simple_vars(
            text: str, data_dict: dict[str, Any], prefix: str = ""
        ) -> str:
            for key, value in data_dict.items():
                var_name = f"{prefix}.{key}" if prefix else key
                placeholder = "{" + var_name + "}"

                if isinstance(value, dict):
                    text = replace_simple_vars(text, value, var_name)
                elif isinstance(value, list):
                    # Simple list handling
                    if all(isinstance(item, str) for item in value):
                        list_content = "\n".join(f"- {item}" for item in value)
                        text = text.replace(placeholder, list_content)
                else:
                    text = text.replace(
                        placeholder, str(value) if value is not None else ""
                    )

            return text

        return replace_simple_vars(content, data)


class DocumentExporter:
    """Document export functionality for multiple formats."""

    def __init__(self):
        self.exporters = {
            "markdown": self._export_markdown,
            "html": self._export_html,
            "pdf": self._export_pdf,
            "docx": self._export_docx,
        }

    def export_document(
        self,
        content: str,
        output_path: str,
        format_type: str,
        style_config: dict[str, Any] | None = None,
    ) -> bool:
        """Export document to specified format."""
        if format_type not in self.exporters:
            raise ValueError(f"Unsupported export format: {format_type}")

        try:
            return self.exporters[format_type](content, output_path, style_config or {})
        except Exception as e:
            logger.error(f"Export failed for {format_type}: {e}")
            return False

    def _export_markdown(
        self, content: str, output_path: str, style_config: dict[str, Any]
    ) -> bool:
        """Export as Markdown file."""
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            return True
        except Exception as e:
            logger.error(f"Markdown export failed: {e}")
            return False

    def _export_html(
        self, content: str, output_path: str, style_config: dict[str, Any]
    ) -> bool:
        """Export as HTML file."""
        try:
            if not markdown:
                raise ImportError("markdown library not available")

            html_content = markdown.markdown(content)

            # Add basic styling
            full_html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Document</title>
    <style>
        body {{
            font-family: {style_config.get("font_family", "Arial, sans-serif")};
            font-size: {style_config.get("font_size", 12)}pt;
            line-height: {style_config.get("line_spacing", 1.4)};
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        h1 {{ color: {style_config.get("colors", {}).get("primary", "#2C3E50")}; }}
        h2 {{ color: {style_config.get("colors", {}).get("secondary", "#34495E")}; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
            """

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(full_html)
            return True
        except Exception as e:
            logger.error(f"HTML export failed: {e}")
            return False

    def _export_pdf(
        self, content: str, output_path: str, style_config: dict[str, Any]
    ) -> bool:
        """Export as PDF file."""
        try:
            if not REPORTLAB_AVAILABLE:
                raise ImportError("reportlab library not available")

            # Create PDF document
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            # Parse markdown-like content and convert to PDF elements
            lines = content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 12))
                elif line.startswith("# "):
                    # Main heading
                    para = Paragraph(line[2:], styles["Title"])
                    story.append(para)
                    story.append(Spacer(1, 12))
                elif line.startswith("## "):
                    # Subheading
                    para = Paragraph(line[3:], styles["Heading2"])
                    story.append(para)
                    story.append(Spacer(1, 6))
                elif line.startswith("### "):
                    # Sub-subheading
                    para = Paragraph(line[4:], styles["Heading3"])
                    story.append(para)
                    story.append(Spacer(1, 6))
                elif line.startswith("- "):
                    # Bullet point
                    para = Paragraph(line, styles["BodyText"])
                    story.append(para)
                elif line.startswith("**") and line.endswith("**"):
                    # Bold text
                    para = Paragraph(f"<b>{line[2:-2]}</b>", styles["BodyText"])
                    story.append(para)
                else:
                    # Regular paragraph
                    para = Paragraph(line, styles["BodyText"])
                    story.append(para)

            doc.build(story)
            return True
        except Exception as e:
            logger.error(f"PDF export failed: {e}")
            return False

    def _export_docx(
        self, content: str, output_path: str, style_config: dict[str, Any]
    ) -> bool:
        """Export as Word document."""
        try:
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx library not available")

            doc = DocxDocument()

            # Parse markdown-like content and convert to Word elements
            lines = content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                elif line.startswith("# "):
                    # Main heading
                    heading = doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    # Subheading
                    heading = doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    # Sub-subheading
                    heading = doc.add_heading(line[4:], level=3)
                elif line.startswith("- "):
                    # Bullet point
                    para = doc.add_paragraph(line[2:], style="List Bullet")
                elif line.startswith("**") and line.endswith("**"):
                    # Bold text
                    para = doc.add_paragraph()
                    run = para.add_run(line[2:-2])
                    run.bold = True
                else:
                    # Regular paragraph
                    doc.add_paragraph(line)

            doc.save(output_path)
            return True
        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
            return False


class DocumentGeneratorPlugin:
    """Main Document Generator Plugin class."""

    def __init__(self):
        self.name = "Document Generator"
        self.version = "1.0.0"
        self.description = "Professional document creation and template system"

        # Initialize components
        self.template_engine = TemplateEngine()
        self.exporter = DocumentExporter()

        # Plugin configuration
        self.config = {
            "output_directory": "generated_documents",
            "default_format": "pdf",
            "auto_save": True,
            "template_directory": "templates",
            "enable_ai_suggestions": True,
        }

        # Ensure output directory exists
        os.makedirs(self.config["output_directory"], exist_ok=True)

    async def initialize(self):
        """Initialize the plugin."""
        logger.info("Document Generator Plugin initialized")

    async def cleanup(self):
        """Cleanup plugin resources."""
        logger.info("Document Generator Plugin cleaned up")

    def capabilities(self) -> list[str]:
        """Return plugin capabilities."""
        return [
            "document_generation",
            "template_management",
            "multi_format_export",
            "resume_builder",
            "meeting_notes",
            "report_generation",
            "markdown_conversion",
        ]

    async def invoke(
        self, action: str, payload: dict[str, Any], context=None
    ) -> dict[str, Any]:
        """Main plugin invocation method."""
        try:
            if action == "list_templates":
                return await self.list_templates(payload.get("category"))
            elif action == "get_template":
                return await self.get_template_details(payload.get("template_id"))
            elif action == "generate_document":
                return await self.generate_document(
                    payload.get("template_id"),
                    payload.get("data"),
                    payload.get("output_format", self.config["default_format"]),
                    payload.get("output_path"),
                )
            elif action == "convert_markdown":
                return await self.convert_markdown(
                    payload.get("content"),
                    payload.get("output_format"),
                    payload.get("output_path"),
                )
            elif action == "create_custom_template":
                return await self.create_custom_template(payload)
            elif action == "get_config":
                return {"status": "success", "data": self.config}
            elif action == "update_config":
                return await self.update_config(payload)
            else:
                return {"status": "error", "message": f"Unknown action: {action}"}
        except Exception as e:
            logger.error(f"Plugin invocation failed: {e}")
            return {"status": "error", "message": str(e)}

    async def list_templates(self, category: str | None = None) -> dict[str, Any]:
        """List available document templates."""
        try:
            templates = self.template_engine.list_templates(category)
            template_data = [asdict(template) for template in templates]

            return {
                "status": "success",
                "data": {
                    "templates": template_data,
                    "count": len(template_data),
                    "categories": list(set(t.category for t in templates)),
                },
            }
        except Exception as e:
            return {"status": "error", "message": f"Failed to list templates: {e}"}

    async def get_template_details(self, template_id: str) -> dict[str, Any]:
        """Get detailed information about a specific template."""
        try:
            template = self.template_engine.get_template(template_id)
            if not template:
                return {
                    "status": "error",
                    "message": f"Template {template_id} not found",
                }

            return {"status": "success", "data": asdict(template)}
        except Exception as e:
            return {
                "status": "error",
                "message": f"Failed to get template details: {e}",
            }

    async def generate_document(
        self,
        template_id: str,
        data: dict[str, Any],
        output_format: str,
        output_path: str | None = None,
    ) -> dict[str, Any]:
        """Generate a document from template and data."""
        try:
            # Get template
            template = self.template_engine.get_template(template_id)
            if not template:
                return {
                    "status": "error",
                    "message": f"Template {template_id} not found",
                }

            # Validate required fields (simplified validation)
            missing_fields = self._validate_template_data(template, data)
            if missing_fields:
                return {
                    "status": "error",
                    "message": f"Missing required fields: {', '.join(missing_fields)}",
                }

            # Render template
            content = self.template_engine.render_template(template_id, data)

            # Generate output path if not provided
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{template_id}_{timestamp}.{output_format}"
                output_path = os.path.join(self.config["output_directory"], filename)

            # Export document
            success = self.exporter.export_document(
                content, output_path, output_format, template.style_config
            )

            if success:
                return {
                    "status": "success",
                    "message": f"Document generated successfully",
                    "data": {
                        "output_path": output_path,
                        "format": output_format,
                        "template_used": template_id,
                    },
                }
            else:
                return {"status": "error", "message": "Document export failed"}

        except Exception as e:
            return {"status": "error", "message": f"Document generation failed: {e}"}

    async def convert_markdown(
        self, content: str, output_format: str, output_path: str
    ) -> dict[str, Any]:
        """Convert markdown content to specified format."""
        try:
            success = self.exporter.export_document(content, output_path, output_format)

            if success:
                return {
                    "status": "success",
                    "message": f"Markdown converted to {output_format}",
                    "data": {"output_path": output_path},
                }
            else:
                return {"status": "error", "message": "Conversion failed"}

        except Exception as e:
            return {"status": "error", "message": f"Markdown conversion failed: {e}"}

    async def create_custom_template(
        self, template_data: dict[str, Any]
    ) -> dict[str, Any]:
        """Create a custom document template."""
        try:
            # Validate template data
            required_fields = ["id", "name", "category", "content_template"]
            missing = [field for field in required_fields if field not in template_data]
            if missing:
                return {
                    "status": "error",
                    "message": f"Missing required fields: {missing}",
                }

            # Create template object
            template = DocumentTemplate(
                id=template_data["id"],
                name=template_data["name"],
                category=template_data["category"],
                description=template_data.get("description", ""),
                fields=template_data.get("fields", []),
                content_template=template_data["content_template"],
                style_config=template_data.get("style_config", {}),
                output_formats=template_data.get(
                    "output_formats", ["pdf", "docx", "html"]
                ),
                created_date=datetime.now().isoformat(),
                author=template_data.get("author", "User"),
            )

            # Store template
            self.template_engine.templates[template.id] = template

            return {
                "status": "success",
                "message": f"Custom template '{template.name}' created",
                "data": {"template_id": template.id},
            }

        except Exception as e:
            return {"status": "error", "message": f"Template creation failed: {e}"}

    async def update_config(self, new_config: dict[str, Any]) -> dict[str, Any]:
        """Update plugin configuration."""
        try:
            self.config.update(new_config)
            return {"status": "success", "message": "Configuration updated"}
        except Exception as e:
            return {"status": "error", "message": f"Config update failed: {e}"}

    def _validate_template_data(
        self, template: DocumentTemplate, data: dict[str, Any]
    ) -> list[str]:
        """Validate that required template fields are present in data."""
        missing_fields = []

        def check_fields(
            fields: list[dict[str, Any]], data_dict: dict[str, Any], prefix: str = ""
        ):
            for field in fields:
                field_name = field["name"]
                full_name = f"{prefix}.{field_name}" if prefix else field_name

                if field.get("required", False) and field_name not in data_dict:
                    missing_fields.append(full_name)
                elif field_name in data_dict and field["type"] == "object":
                    if "fields" in field:
                        check_fields(field["fields"], data_dict[field_name], full_name)

        check_fields(template.fields, data)
        return missing_fields


# Plugin entry point
def get_plugin():
    """Return the plugin instance."""
    return DocumentGeneratorPlugin()


# For testing
if __name__ == "__main__":
    import asyncio

    async def test_plugin():
        plugin = DocumentGeneratorPlugin()
        await plugin.initialize()

        # Test template listing
        result = await plugin.list_templates()
        print("Templates:", result)

        # Test resume generation
        resume_data = {
            "personal_info": {
                "full_name": "John Doe",
                "email": "john.doe@example.com",
                "phone": "+1 (555) 123-4567",
                "location": "San Francisco, CA",
                "website": "https://johndoe.dev",
                "linkedin": "https://linkedin.com/in/johndoe",
            },
            "professional_summary": "Experienced software engineer with 5+ years in full-stack development...",
            "experience": [
                {
                    "company": "Tech Corp",
                    "position": "Senior Software Engineer",
                    "start_date": "2020-01",
                    "end_date": None,
                    "description": "Led development of microservices architecture...",
                    "achievements": [
                        "Improved system performance by 40%",
                        "Led team of 5 developers",
                    ],
                }
            ],
            "education": [
                {
                    "institution": "University of Technology",
                    "degree": "Bachelor of Science",
                    "field": "Computer Science",
                    "graduation_date": "2019-05",
                    "gpa": 3.8,
                }
            ],
            "skills": ["Python", "JavaScript", "React", "Node.js", "AWS"],
            "certifications": ["AWS Certified Developer"],
        }

        result = await plugin.generate_document(
            "professional_resume", resume_data, "pdf"
        )
        print("Resume generation:", result)

        await plugin.cleanup()

    asyncio.run(test_plugin())
